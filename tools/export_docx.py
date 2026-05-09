"""
export_docx.py — 导出 Word 文档工具
将合并后的 novel_full.md 转换为格式化的 .docx 文件。
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def md_to_docx(md_path: Path, out_path: Path) -> None:
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    content = md_path.read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    lines = content.split("\n")
    for line in lines:
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            line_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            line_clean = re.sub(r"\*(.+?)\*", r"\1", line_clean)
            p = doc.add_paragraph(line_clean)
            p.paragraph_format.first_line_indent = Pt(24)

    doc.save(str(out_path))


def run_export_docx(project_root: Path) -> str:
    export_dir = project_root / "project_repo/manuscript/final_export"
    md_path = export_dir / "novel_full.md"

    if not md_path.exists():
        return "❌ novel_full.md 不存在，请先运行 merge_manuscript。"

    out_path = export_dir / "novel_full.docx"

    try:
        md_to_docx(md_path, out_path)
        size_kb = out_path.stat().st_size // 1024
        return f"✅ DOCX 导出完成\n- 输出文件: {out_path}\n- 文件大小: {size_kb} KB"
    except ImportError as e:
        return f"❌ {e}"
    except Exception as e:
        return f"❌ 导出失败: {e}"


if __name__ == "__main__":
    root = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(".")
    print(run_export_docx(root))
